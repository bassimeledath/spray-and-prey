from docx import Document
from docx.shared import Pt
from models import Model
from constants import TAGLINE_INDEX, SKILLS_INDEX, OUTPUT_RESUME_PATH, FONT_SIZE
from prompts import prompt_tag_section, prompt_skills_section


class Prey:
    def __init__(self, doc_path, context):
        self.context = context
        self.doc = Document(doc_path)
        self.model = Model()
        self.new_skills = ""
        self.new_tag = ""

    def generate_tag_section(self):
        prompt = prompt_tag_section(context=self.context)
        self.new_tag = self.model.generate_text(prompt)

    def generate_skills_section(self):
        prompt = prompt_skills_section(context=self.context)
        self.new_skills = self.model.generate_text(prompt)

    def get_tag_and_skills(self):
        self.generate_skills_section()
        self.generate_tag_section()
        return self.new_tag, self.new_skills

    def modify_resume(self, new_tag, new_skills):
        if not new_tag or not new_skills:
            raise ValueError("Please generate tag and skills first")
        try:
            font_size = Pt(FONT_SIZE)

            # Modify the tagline
            tagline_paragraph = self.doc.paragraphs[TAGLINE_INDEX]
            tagline_paragraph.clear()
            run = tagline_paragraph.add_run(new_tag)
            run.font.size = font_size

            # Modify the skills
            skills_paragraph = self.doc.paragraphs[SKILLS_INDEX]
            skills_paragraph.clear()
            run = skills_paragraph.add_run(new_skills)
            run.font.size = font_size

        except Exception as e:
            print(e)

        self.doc.save(OUTPUT_RESUME_PATH)
